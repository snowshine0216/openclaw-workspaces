#!/usr/bin/env python3
"""
Generate daily vocabulary review .docx
Usage: python generate_review.py [--db-path PATH]
"""
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

from vocab_manager import select_for_review, update_review_status, get_active_count


def generate_docx(items):
    """Generate a .docx file with vocabulary review content."""
    doc = Document()

    # Title
    title = doc.add_heading('📚 Daily Vocabulary Review', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Spacer

    if not items:
        doc.add_paragraph("No vocabulary items to review today. Great job! 🎉")
        return doc

    for i, item in enumerate(items, 1):
        content = item.get('content', '')

        # Section header
        doc.add_heading(f"{i}. {content}", level=1)

        # IPA pronunciation
        ipa = item.get('ipa', '')
        if ipa:
            p = doc.add_paragraph()
            p.add_run('IPA: ').bold = True
            p.add_run(ipa)

        # English definition
        english = item.get('english', '')
        if english:
            p = doc.add_paragraph()
            p.add_run('English: ').bold = True
            p.add_run(english)

        # Chinese definition
        chinese = item.get('chinese', '')
        if chinese:
            p = doc.add_paragraph()
            p.add_run('中文: ').bold = True
            p.add_run(chinese)

        # Example
        example = item.get('example', '')
        if example:
            p = doc.add_paragraph()
            p.add_run('Example: ').bold = True
            p.add_run(example).italic = True

        # Synonyms
        synonyms = item.get('synonyms', '[]')
        if isinstance(synonyms, str):
            import json
            synonyms = json.loads(synonyms)
        if synonyms:
            p = doc.add_paragraph()
            p.add_run('Synonyms: ').bold = True
            p.add_run(', '.join(synonyms))

        # Original context
        original_context = item.get('original_context', '')
        if original_context:
            p = doc.add_paragraph()
            p.add_run('📝 Original Context: ').bold = True
            p.add_run(original_context).italic = True

        # Fun fact
        fun_fact = item.get('fun_fact', '')
        if fun_fact:
            p = doc.add_paragraph()
            p.add_run('🎯 Fun Fact: ').bold = True
            p.add_run(fun_fact)

        # Memory trick
        memory_trick = item.get('memory_trick', '')
        if memory_trick:
            p = doc.add_paragraph()
            p.add_run('💡 Memory Trick: ').bold = True
            p.add_run(memory_trick)

        # Review status
        review_count = item.get('review_count', 0)
        status = item.get('status', 'learning')
        p = doc.add_paragraph()
        p.add_run(f"Review #: {review_count} | Status: {status}").italic = True

        doc.add_paragraph()  # Spacer

    return doc


def main():
    parser = argparse.ArgumentParser(description='Generate daily vocabulary review')
    parser.add_argument('--db-path', type=str,
                       default='../../projects/vocabulary/data/vocab.db',
                       help='Path to vocabulary database')
    parser.add_argument('--output-dir', type=str,
                       default='../../projects/vocabulary/output',
                       help='Path to output directory')
    args = parser.parse_args()

    # Resolve paths relative to script location
    script_dir = Path(__file__).parent
    db_path = (script_dir / args.db_path).resolve()
    output_dir = (script_dir / args.output_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 50)
    print("📚 Vocabulary Review Generator")
    print("=" * 50)

    # Select items for review
    items = select_for_review(str(db_path), max_items=20)
    print(f"[INFO] Selected {len(items)} items for review")
    print(f"[INFO] Active vocabulary: {get_active_count(str(db_path))}")

    # Generate .docx
    doc = generate_docx(items)
    output_file = output_dir / "daily_review.docx"
    doc.save(str(output_file))
    print(f"[INFO] Generated: {output_file}")

    # Update review status
    reviewed_ids = [item['id'] for item in items]
    update_review_status(str(db_path), reviewed_ids)
    print(f"[INFO] Updated review status for {len(items)} items")

    print("=" * 50)
    print("✅ Daily review document generated!")
    print("=" * 50)


if __name__ == '__main__':
    main()
