#!/usr/bin/env python3
"""
Vocabulary Review System for Snow's English Learning
Generates daily .docx review and sends to Feishu
"""

import json
import os
import sys
import base64
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Paths
SCRIPT_DIR = Path(__file__).parent
VOCAB_FILE = SCRIPT_DIR / "vocab.json"
OUTPUT_FILE = SCRIPT_DIR / "daily_review.docx"

# Feishu Chat ID for this agent
FEISHU_CHAT_ID = "oc_55bf80b97398600ff6da478ae62937de"


def load_vocab():
    """Load vocabulary data from JSON file."""
    with open(VOCAB_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def save_vocab(data):
    """Save vocabulary data to JSON file."""
    with open(VOCAB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def select_items_for_review(data):
    """Select items for daily review using spaced repetition."""
    items = data.get("items", [])
    config = data.get("config", {})
    max_items = config.get("items_per_review", 20)

    if not items:
        return []

    # Sort by: learning status first, then by review count (ascending), then by oldest review
    def sort_key(item):
        status = item.get("status", "learning")
        review_count = item.get("review_count", 0)
        last_reviewed = item.get("last_reviewed", "")

        # Priority: learning > reviewing > mastered
        status_priority = {"learning": 0, "reviewing": 1, "mastered": 2}.get(status, 0)

        return (status_priority, review_count, last_reviewed or "")

    sorted_items = sorted(items, key=sort_key)

    # Select top items up to max_items
    selected = sorted_items[:max_items]

    return selected


def generate_docx(items):
    """Generate a .docx file with vocabulary review content."""
    doc = Document()

    # Title
    title = doc.add_heading("📚 Daily Vocabulary Review", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()  # Spacer

    if not items:
        doc.add_paragraph("No vocabulary items to review today. Great job! 🎉")
        return doc

    for i, item in enumerate(items, 1):
        item_type = item.get("type", "word")
        content = item.get("content", "")

        # Section header
        doc.add_heading(f"{i}. {content}", level=1)

        # IPA pronunciation
        ipa = item.get("ipa", "")
        if ipa:
            p = doc.add_paragraph()
            p.add_run("IPA: ").bold = True
            p.add_run(ipa)

        if item_type == "word" or item_type == "phrase":
            # English definition
            english = item.get("english", "")
            if english:
                p = doc.add_paragraph()
                p.add_run("English: ").bold = True
                p.add_run(english)

            # Chinese definition
            chinese = item.get("chinese", "")
            if chinese:
                p = doc.add_paragraph()
                p.add_run("中文: ").bold = True
                p.add_run(chinese)

            # Example
            example = item.get("example", "")
            if example:
                p = doc.add_paragraph()
                p.add_run("Example: ").bold = True
                p.add_run(example).italic = True

            # Synonyms
            synonyms = item.get("synonyms", [])
            if synonyms:
                p = doc.add_paragraph()
                p.add_run("Synonyms: ").bold = True
                p.add_run(", ".join(synonyms))

            # Original context (optional)
            original_context = item.get("original_context", "")
            if original_context:
                p = doc.add_paragraph()
                p.add_run("📝 Original Context: ").bold = True
                p.add_run(original_context).italic = True

            # Fun fact (optional)
            fun_fact = item.get("fun_fact", "")
            if fun_fact:
                p = doc.add_paragraph()
                p.add_run("🎯 Fun Fact: ").bold = True
                p.add_run(fun_fact)

            # Memory trick
            memory_trick = item.get("memory_trick", "")
            if memory_trick:
                p = doc.add_paragraph()
                p.add_run("💡 Memory Trick: ").bold = True
                p.add_run(memory_trick)

        else:  # sentence/paragraph
            # Chinese translation
            chinese = item.get("chinese", "")
            if chinese:
                p = doc.add_paragraph()
                p.add_run("中文翻译: ").bold = True
                p.add_run(chinese)

            # Context
            context = item.get("context", "")
            if context:
                p = doc.add_paragraph()
                p.add_run("When to use: ").bold = True
                p.add_run(context)

            # Alternatives
            alternatives = item.get("alternatives", [])
            if alternatives:
                p = doc.add_paragraph()
                p.add_run("Alternatives: ").bold = True
                p.add_run(", ".join(alternatives))

            # Key phrases
            key_phrases = item.get("key_phrases", [])
            if key_phrases:
                p = doc.add_paragraph()
                p.add_run("Key Phrases: ").bold = True
                p.add_run(", ".join(key_phrases))

        # Review status
        review_count = item.get("review_count", 0)
        status = item.get("status", "learning")
        p = doc.add_paragraph()
        p.add_run(f"Review #: {review_count} | Status: {status}").italic = True

        doc.add_paragraph()  # Spacer between items

    return doc


def update_review_status(data, selected_items):
    """Update review count and status for reviewed items."""
    today = datetime.now().strftime("%Y-%m-%d")

    for selected in selected_items:
        for item in data["items"]:
            if item.get("id") == selected.get("id"):
                item["review_count"] = item.get("review_count", 0) + 1
                item["last_reviewed"] = today

                # Update status based on review count
                count = item["review_count"]
                if count >= 3:
                    item["status"] = "reviewing"
                if count >= 7:
                    item["status"] = "mastered"
                break

    save_vocab(data)


def remove_old_file():
    """Remove old daily_review.docx if it exists."""
    if OUTPUT_FILE.exists():
        OUTPUT_FILE.unlink()
        print(f"[INFO] Removed old file: {OUTPUT_FILE}")


def main():
    print("=" * 50)
    print("📚 Vocabulary Review System")
    print("=" * 50)

    # Remove old output file first
    remove_old_file()

    # Load vocabulary
    data = load_vocab()
    print(f"[INFO] Loaded {len(data.get('items', []))} vocabulary items")

    # Select items for review
    selected = select_items_for_review(data)
    print(f"[INFO] Selected {len(selected)} items for today's review")

    # Generate .docx
    doc = generate_docx(selected)
    doc.save(str(OUTPUT_FILE))
    print(f"[INFO] Generated: {OUTPUT_FILE}")

    # Update review status
    update_review_status(data, selected)
    print(f"[INFO] Updated review status for {len(selected)} items")

    print("=" * 50)
    print("✅ Daily review document generated!")
    print(f"📄 File: {OUTPUT_FILE}")
    print(f"🎯 Feishu Chat ID: {FEISHU_CHAT_ID}")
    print("=" * 50)

    # Return info for external sending
    return {
        "file_path": str(OUTPUT_FILE),
        "chat_id": FEISHU_CHAT_ID,
        "items_count": len(selected)
    }


if __name__ == "__main__":
    result = main()
    
    # Output JSON for easy parsing by caller
    print("\n--- RESULT ---")
    print(json.dumps(result))
